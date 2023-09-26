from docx import Document
import genanki
import random

deck_id = random.randrange(1 << 30, 1 << 31)
print(deck_id)

deck_name = input('Deck Name: ')
deck_name += '.apkg'

file_name = input('Docx file name: ')

if file_name.find('.docx') == -1:
    file_name += '.docx'

my_deck = genanki.Deck(int(deck_id), deck_name)

doc = Document(file_name)

deck = genanki.Deck(deck_id, deck_name)

save = False

title = ""
info = ""
titles = []
infos = []
list_paragraph = False

for paragraph in doc.paragraphs:

    if paragraph.style.name == 'Heading 1':
        continue

    if paragraph.style.name == 'Heading 2':
        if info == '':
            continue

        infos.append(info)
        info = ''
        continue

    if paragraph.style.name == 'List Paragraph':
        if list_paragraph:
            info += '\n<li>' + paragraph.text + '</li>\n'
        else:
            info += '\n <br><ul><li>' + paragraph.text + '</li>\n'
            list_paragraph = True

    if paragraph.style.name == 'Normal':
        if not list_paragraph:
            info += paragraph.text + '\n'
        else:
            info += '</ul><br>' + paragraph.text + '\n'
            list_paragraph = False

    # print(paragraph.style.name)

infos.append(info)

# print(infos[0])

for paragraph in doc.paragraphs:

    if paragraph.style.name == 'Heading 2':
        titles.append(paragraph.text)

    # print(paragraph.style.name)

print(str(len(titles)) + ' headings found')
print(str(len(infos)) + ' texts found')

if len(infos) != len(titles):
    print("Failed, mismatch.")
    exit(0)

c = 0
card = 'Card '
full_name = ''

models = [genanki.Model(
    deck_id,
    'Simple Model',
    fields=[
        {'name': 'Question'},
        {'name': 'Answer'},
    ],
    templates=[
        {
            'name': 'Card 1',
            'qfmt': '{{Question}}',
            'afmt': '{{FrontSide}}<hr id="answer">{{Answer}}',
        },
    ]) for i in range(len(infos))]

for i in range(len(infos)):
    full_name = card + str(c + 1)
    models[c].templates[0]['name'] = full_name

    note = genanki.Note(
        model=models[c],
        fields=[titles[i], infos[i]])
    my_deck.add_note(note)
    c += 1

genanki.Package(my_deck).write_to_file(deck_name)
print('Finished')
